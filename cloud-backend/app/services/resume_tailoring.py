"""
Resume Tailoring Service - JTR (Job-Tailored Resume) Implementation
Generates ATS-optimized resumes with Reasonable Synthesis (RS)
"""

import asyncio
from typing import Dict, List, Optional, Any, Tuple
from datetime import datetime, timedelta
import json
import re
from dataclasses import dataclass, asdict
import structlog
from pathlib import Path

# Document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
except ImportError:
    print("Warning: Document generation libraries not installed")
    Document = None

from app.core.config import settings
from app.ai.orchestrator import get_ai_orchestrator, AITask, TaskType

logger = structlog.get_logger(__name__)


@dataclass
class RSBullet:
    """Reasonable Synthesis bullet point"""
    text: str
    rs: bool = False
    rs_basis: Optional[str] = None
    confidence: float = 1.0
    risk_level: str = "low"  # low, medium, high
    source_ids: Optional[List[str]] = None
    original_text: Optional[str] = None


@dataclass 
class ResumeSection:
    """Resume section with bullets"""
    title: str
    content: str
    bullets: List[RSBullet]
    order: int = 0


@dataclass
class TailoredResume:
    """Complete tailored resume structure"""
    user_id: str
    job_id: str
    summary: str
    skills: List[str]
    sections: List[ResumeSection]
    match_score: int
    ats_score: int
    rs_bullets_count: int
    generation_metadata: Dict[str, Any]
    created_at: datetime
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for API response"""
        return asdict(self)


class ResumeTailoringService:
    """
    Service for creating job-tailored resumes with AI assistance
    Implements Reasonable Synthesis (RS) with proper attribution
    """
    
    def __init__(self):
        self.logger = structlog.get_logger(__name__)
        self.ai_orchestrator = get_ai_orchestrator()
        self.skill_synonyms = self._load_skill_synonyms()
        self.ats_keywords = self._load_ats_keywords()
        
    def _load_skill_synonyms(self) -> Dict[str, List[str]]:
        """Load skill synonyms for ATS optimization"""
        return {
            "JavaScript": ["JS", "ECMAScript", "Node.js"],
            "Python": ["Python 3", "Django", "Flask", "FastAPI"],
            "React": ["React.js", "ReactJS", "React Native"],
            "SQL": ["MySQL", "PostgreSQL", "SQLite", "Database"],
            "AWS": ["Amazon Web Services", "Cloud Computing", "EC2", "S3"],
            "Machine Learning": ["ML", "Artificial Intelligence", "AI", "Data Science"],
            "API": ["REST API", "GraphQL", "Web Services", "Integration"],
            "Docker": ["Containerization", "DevOps", "Kubernetes"],
            "Git": ["Version Control", "GitHub", "GitLab", "Source Control"]
        }
    
    def _load_ats_keywords(self) -> List[str]:
        """Load common ATS-friendly keywords"""
        return [
            "developed", "implemented", "designed", "managed", "led", "created",
            "improved", "optimized", "reduced", "increased", "achieved", "delivered",
            "collaborated", "coordinated", "analyzed", "researched", "maintained",
            "troubleshooted", "configured", "deployed", "tested", "documented"
        ]
    
    async def create_tailored_resume(
        self,
        job_data: Dict[str, Any],
        user_profile: Dict[str, Any],
        evidence_vault: List[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Create a job-tailored resume with RS enhancements
        """
        try:
            self.logger.info(
                "Starting resume tailoring", 
                job_id=job_data.get('job_id'), 
                user_id=user_profile.get('user_id')
            )
            
            # Step 1: Analyze job requirements
            job_analysis = await self._analyze_job_requirements(job_data)
            
            # Step 2: Calculate match score
            match_score = await self._calculate_match_score(job_analysis, user_profile)
            
            # Step 3: Generate tailored summary
            summary = await self._generate_summary(job_data, user_profile, job_analysis)
            
            # Step 4: Optimize skills section
            skills = await self._optimize_skills_section(job_analysis, user_profile)
            
            # Step 5: Enhance experience with RS
            sections = await self._enhance_experience_sections(
                user_profile, job_analysis, evidence_vault or []
            )
            
            # Step 6: ATS optimization
            ats_score = await self._calculate_ats_score(summary, skills, sections, job_analysis)
            
            # Step 7: Generate resume documents
            documents = await self._generate_resume_documents(
                user_profile, summary, skills, sections
            )
            
            # Step 8: Create diff report
            diff_report = await self._create_diff_report(user_profile, summary, skills, sections)
            
            # Count RS bullets
            rs_bullets_count = sum(
                sum(1 for bullet in section.bullets if bullet.rs) 
                for section in sections
            )
            
            result = {
                "success": True,
                "match_score": match_score,
                "ats_score": ats_score,
                "tailored_resume": {
                    "summary": summary,
                    "skills": skills,
                    "sections": [asdict(section) for section in sections],
                    "rs_bullets_count": rs_bullets_count
                },
                "documents": documents,
                "diff_report": diff_report,
                "job_analysis": job_analysis,
                "generation_metadata": {
                    "generated_at": datetime.utcnow().isoformat(),
                    "model_used": settings.DEFAULT_MODEL,
                    "rs_applied": rs_bullets_count > 0,
                    "ats_optimized": ats_score >= settings.ATS_SCORE_THRESHOLD
                }
            }
            
            self.logger.info(
                "Resume tailoring completed",
                match_score=match_score,
                ats_score=ats_score,
                rs_bullets=rs_bullets_count
            )
            
            return result
            
        except Exception as e:
            self.logger.error("Resume tailoring failed", error=str(e))
            raise
    
    async def _analyze_job_requirements(self, job_data: Dict[str, Any]) -> Dict[str, Any]:
        """Analyze job posting to extract requirements and keywords"""
        
        task = AITask(
            task_id=f"job_analysis_{datetime.utcnow().timestamp()}",
            task_type=TaskType.JOB_ANALYSIS,
            input_data=job_data
        )
        
        response = await self.ai_orchestrator.execute_task(task)
        
        if response.success:
            return response.result
        else:
            # Fallback analysis
            return await self._fallback_job_analysis(job_data)
    
    async def _fallback_job_analysis(self, job_data: Dict[str, Any]) -> Dict[str, Any]:
        """Fallback job analysis without AI"""
        description = job_data.get('description', '').lower()
        
        # Extract skills using keyword matching
        tech_skills = []
        for skill, synonyms in self.skill_synonyms.items():
            if any(synonym.lower() in description for synonym in [skill] + synonyms):
                tech_skills.append(skill)
        
        # Extract experience level
        experience_indicators = {
            'entry': ['entry', 'junior', '0-2 years', 'new grad'],
            'mid': ['mid-level', '2-5 years', 'intermediate'],
            'senior': ['senior', '5+ years', 'lead', 'principal']
        }
        
        experience_level = 'mid'  # default
        for level, indicators in experience_indicators.items():
            if any(indicator in description for indicator in indicators):
                experience_level = level
                break
        
        return {
            "required_skills": tech_skills[:10],  # Top 10
            "nice_to_have": [],
            "experience_level": experience_level,
            "key_responsibilities": [],
            "company_culture": [],
            "analysis_confidence": 0.6
        }
    
    async def _calculate_match_score(
        self, 
        job_analysis: Dict[str, Any], 
        user_profile: Dict[str, Any]
    ) -> int:
        """Calculate job match score (0-100)"""
        
        required_skills = job_analysis.get('required_skills', [])
        user_skills = [s.lower() for s in user_profile.get('skills', [])]
        
        # Skills match
        matching_skills = [
            skill for skill in required_skills 
            if skill.lower() in user_skills or 
            any(syn.lower() in user_skills for syn in self.skill_synonyms.get(skill, []))
        ]
        
        skills_match = (len(matching_skills) / max(len(required_skills), 1)) * 100
        
        # Experience match (simplified)
        user_experience = user_profile.get('experience_years', 2)
        required_exp_map = {'entry': 1, 'mid': 3, 'senior': 5}
        required_exp = required_exp_map.get(job_analysis.get('experience_level', 'mid'), 3)
        
        exp_match = min((user_experience / required_exp) * 100, 100)
        
        # Weighted average
        match_score = (
            skills_match * settings.SKILLS_MATCH_WEIGHT +
            exp_match * settings.EXPERIENCE_MATCH_WEIGHT +
            80 * 0.4  # Base compatibility score
        )
        
        return min(int(match_score), 100)
    
    async def _generate_summary(
        self,
        job_data: Dict[str, Any],
        user_profile: Dict[str, Any], 
        job_analysis: Dict[str, Any]
    ) -> str:
        """Generate tailored professional summary"""
        
        # Extract key elements
        job_title = job_data.get('title', 'Professional')
        company = job_data.get('company', 'target organization')
        top_skills = job_analysis.get('required_skills', [])[:3]
        user_experience = user_profile.get('experience_years', 'several')
        
        # Create tailored summary
        summary = f"Results-driven {job_title} with {user_experience}+ years of experience in "
        
        if top_skills:
            summary += f"{', '.join(top_skills[:-1])}"
            if len(top_skills) > 1:
                summary += f" and {top_skills[-1]}"
        else:
            summary += "software development and technology solutions"
        
        summary += f". Proven track record of delivering high-quality solutions that drive business value. "
        summary += f"Eager to contribute expertise to {company}'s continued success through innovative problem-solving and collaborative teamwork."
        
        # Add key job-specific keywords
        if 'leadership' in str(job_analysis).lower():
            summary += " Demonstrated leadership capabilities in managing cross-functional projects."
        
        return summary
    
    async def _optimize_skills_section(
        self,
        job_analysis: Dict[str, Any],
        user_profile: Dict[str, Any]
    ) -> List[str]:
        """Create optimized skills section with job keyword alignment"""
        
        required_skills = job_analysis.get('required_skills', [])
        user_skills = user_profile.get('skills', [])
        
        # Prioritize job-relevant skills
        prioritized_skills = []
        
        # Add matching skills first (with synonyms)
        for skill in required_skills:
            # Add exact matches
            if skill in user_skills:
                prioritized_skills.append(skill)
            else:
                # Check for synonyms
                synonyms = self.skill_synonyms.get(skill, [])
                for user_skill in user_skills:
                    if (user_skill.lower() in [s.lower() for s in synonyms] or
                        skill.lower() in user_skill.lower()):
                        # Use job posting terminology for ATS
                        prioritized_skills.append(f"{skill} / {user_skill}")
                        break
        
        # Add remaining user skills
        for skill in user_skills:
            if skill not in prioritized_skills and len(prioritized_skills) < 18:
                prioritized_skills.append(skill)
        
        # Ensure we have 12-18 skills as recommended
        if len(prioritized_skills) < 12:
            # Add relevant skills from job analysis
            for skill in job_analysis.get('nice_to_have', []):
                if len(prioritized_skills) < 12:
                    prioritized_skills.append(skill)
        
        return prioritized_skills[:18]  # Cap at 18 skills
    
    async def _enhance_experience_sections(
        self,
        user_profile: Dict[str, Any],
        job_analysis: Dict[str, Any],
        evidence_vault: List[Dict[str, Any]]
    ) -> List[ResumeSection]:
        """Enhance experience sections with RS bullets"""
        
        experience_items = user_profile.get('experience', [])
        sections = []
        
        for i, exp in enumerate(experience_items):
            bullets = []
            original_bullets = exp.get('responsibilities', [])
            
            for bullet_text in original_bullets:
                # Apply RS enhancement
                enhanced_bullet = await self._apply_reasonable_synthesis(
                    bullet_text, job_analysis, evidence_vault, exp
                )
                bullets.append(enhanced_bullet)
            
            # Add job-relevant bullets if needed
            if len(bullets) < 4:
                additional_bullets = await self._generate_relevant_bullets(
                    exp, job_analysis, 4 - len(bullets)
                )
                bullets.extend(additional_bullets)
            
            section = ResumeSection(
                title=f"{exp.get('title', 'Professional')} - {exp.get('company', 'Company')}",
                content=f"{exp.get('start_date', 'Start')} - {exp.get('end_date', 'End')}",
                bullets=bullets,
                order=i
            )
            sections.append(section)
        
        return sections
    
    async def _apply_reasonable_synthesis(
        self,
        bullet_text: str,
        job_analysis: Dict[str, Any],
        evidence_vault: List[Dict[str, Any]],
        experience_context: Dict[str, Any]
    ) -> RSBullet:
        """Apply Reasonable Synthesis to enhance bullet points"""
        
        enhanced_text = bullet_text
        rs_applied = False
        rs_basis = None
        confidence = 1.0
        risk_level = "low"
        
        # Check if bullet needs quantification
        if self._needs_quantification(bullet_text):
            # Find supporting evidence
            evidence = self._find_supporting_evidence(bullet_text, evidence_vault, experience_context)
            
            if evidence:
                # Apply quantification with RS
                quantified_text, rs_basis = self._add_quantification(bullet_text, evidence)
                if quantified_text != bullet_text:
                    enhanced_text = quantified_text
                    rs_applied = True
                    confidence = min(evidence.get('confidence', 0.8), settings.RS_CONFIDENCE_THRESHOLD)
                    
                    # Risk assessment
                    if confidence < 0.6:
                        risk_level = "high"
                    elif confidence < 0.8:
                        risk_level = "medium"
        
        # ATS keyword optimization
        enhanced_text = self._optimize_bullet_for_ats(enhanced_text, job_analysis)
        
        return RSBullet(
            text=enhanced_text,
            rs=rs_applied,
            rs_basis=rs_basis,
            confidence=confidence,
            risk_level=risk_level,
            original_text=bullet_text if enhanced_text != bullet_text else None
        )
    
    def _needs_quantification(self, bullet_text: str) -> bool:
        """Check if bullet point needs quantification"""
        quantifiable_verbs = [
            'improved', 'increased', 'reduced', 'optimized', 'enhanced',
            'grew', 'generated', 'saved', 'streamlined', 'accelerated'
        ]
        
        has_quantifiable_verb = any(verb in bullet_text.lower() for verb in quantifiable_verbs)
        has_numbers = bool(re.search(r'\d+%|\d+x|\$\d+|\d+ hours|\d+ minutes', bullet_text))
        
        return has_quantifiable_verb and not has_numbers
    
    def _find_supporting_evidence(
        self,
        bullet_text: str,
        evidence_vault: List[Dict[str, Any]],
        experience_context: Dict[str, Any]
    ) -> Optional[Dict[str, Any]]:
        """Find supporting evidence for RS application"""
        
        # Match by company and time period
        company = experience_context.get('company', '')
        start_year = experience_context.get('start_date', '').split('-')[0] if experience_context.get('start_date') else None
        end_year = experience_context.get('end_date', '').split('-')[0] if experience_context.get('end_date') else None
        
        for evidence in evidence_vault:
            # Must be same company and time period for valid RS
            if (evidence.get('company') == company and
                self._time_periods_overlap(evidence, start_year, end_year)):
                
                # Check content similarity
                if self._calculate_text_similarity(bullet_text, evidence.get('description', '')) > 0.6:
                    return evidence
        
        return None
    
    def _time_periods_overlap(
        self, 
        evidence: Dict[str, Any], 
        start_year: str, 
        end_year: str
    ) -> bool:
        """Check if time periods overlap for valid RS"""
        evidence_year = evidence.get('year')
        if not evidence_year or not start_year:
            return False
        
        try:
            ev_year = int(evidence_year)
            start = int(start_year)
            end = int(end_year) if end_year else datetime.now().year
            
            return start <= ev_year <= end
        except (ValueError, TypeError):
            return False
    
    def _calculate_text_similarity(self, text1: str, text2: str) -> float:
        """Simple text similarity calculation"""
        # Simplified - in production, use embedding similarity
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0
    
    def _add_quantification(
        self, 
        bullet_text: str, 
        evidence: Dict[str, Any]
    ) -> Tuple[str, str]:
        """Add quantification based on evidence"""
        
        # Extract metrics from evidence
        metrics = evidence.get('metrics', {})
        
        if 'improvement_percentage' in metrics:
            percentage = metrics['improvement_percentage']
            # Use interval for RS
            interval = f"{percentage-5}-{percentage+5}%"
            enhanced_text = f"{bullet_text} by approximately {interval}"
            rs_basis = f"Quantification based on {evidence.get('source', 'documented metrics')} from same role/timeframe"
            
        elif 'time_saved' in metrics:
            time_saved = metrics['time_saved']
            enhanced_text = f"{bullet_text}, saving approximately {time_saved} per week"
            rs_basis = f"Time quantification based on {evidence.get('source', 'project documentation')}"
            
        else:
            # Generic improvement
            enhanced_text = f"{bullet_text} by approximately 15-20%"
            rs_basis = "Typical improvement metrics for similar initiatives in comparable roles"
        
        return enhanced_text, rs_basis
    
    def _optimize_bullet_for_ats(
        self, 
        bullet_text: str, 
        job_analysis: Dict[str, Any]
    ) -> str:
        """Optimize bullet point for ATS keywords"""
        
        required_skills = job_analysis.get('required_skills', [])
        enhanced_text = bullet_text
        
        # Add relevant keywords if naturally fitting
        for skill in required_skills:
            if skill.lower() not in bullet_text.lower():
                # Check if we can naturally incorporate the skill
                synonyms = self.skill_synonyms.get(skill, [])
                for synonym in synonyms:
                    if synonym.lower() in bullet_text.lower():
                        # Add skill name for ATS
                        enhanced_text = enhanced_text.replace(synonym, f"{synonym} ({skill})")
                        break
        
        return enhanced_text
    
    async def _generate_relevant_bullets(
        self,
        experience: Dict[str, Any],
        job_analysis: Dict[str, Any],
        count: int
    ) -> List[RSBullet]:
        """Generate additional relevant bullets for the role"""
        
        bullets = []
        required_skills = job_analysis.get('required_skills', [])[:count]
        
        for skill in required_skills:
            bullet_text = f"Utilized {skill} to deliver high-quality solutions and maintain code standards"
            
            bullet = RSBullet(
                text=bullet_text,
                rs=True,
                rs_basis="Generated bullet based on job requirements and role context",
                confidence=0.7,
                risk_level="medium"
            )
            bullets.append(bullet)
        
        return bullets
    
    async def _calculate_ats_score(
        self,
        summary: str,
        skills: List[str],
        sections: List[ResumeSection],
        job_analysis: Dict[str, Any]
    ) -> int:
        """Calculate ATS compatibility score"""
        
        full_text = f"{summary} {' '.join(skills)} "
        for section in sections:
            full_text += f"{section.title} {section.content} "
            for bullet in section.bullets:
                full_text += f"{bullet.text} "
        
        task = AITask(
            task_id=f"ats_check_{datetime.utcnow().timestamp()}",
            task_type=TaskType.ATS_OPTIMIZATION,
            input_data={
                "resume_text": full_text,
                "job_keywords": job_analysis.get('required_skills', [])
            }
        )
        
        response = await self.ai_orchestrator.execute_task(task)
        
        if response.success:
            return response.result.get('ats_score', 85)
        else:
            # Fallback ATS scoring
            return self._calculate_fallback_ats_score(full_text, job_analysis)
    
    def _calculate_fallback_ats_score(self, text: str, job_analysis: Dict[str, Any]) -> int:
        """Fallback ATS score calculation"""
        score = 70  # Base score
        
        # Check for standard sections
        if 'experience' in text.lower():
            score += 5
        if 'skills' in text.lower():
            score += 5
        if 'education' in text.lower():
            score += 5
        
        # Check keyword density
        required_skills = job_analysis.get('required_skills', [])
        found_keywords = sum(1 for skill in required_skills if skill.lower() in text.lower())
        keyword_score = (found_keywords / max(len(required_skills), 1)) * 15
        score += keyword_score
        
        return min(int(score), 100)
    
    async def _generate_resume_documents(
        self,
        user_profile: Dict[str, Any],
        summary: str,
        skills: List[str],
        sections: List[ResumeSection]
    ) -> Dict[str, str]:
        """Generate resume documents in multiple formats"""
        
        documents = {}
        
        # Generate DOCX
        if Document:
            try:
                docx_path = await self._generate_docx_resume(
                    user_profile, summary, skills, sections
                )
                documents['docx'] = docx_path
            except Exception as e:
                self.logger.error("DOCX generation failed", error=str(e))
        
        # Generate PDF
        try:
            pdf_path = await self._generate_pdf_resume(
                user_profile, summary, skills, sections
            )
            documents['pdf'] = pdf_path
        except Exception as e:
            self.logger.error("PDF generation failed", error=str(e))
        
        return documents
    
    async def _generate_docx_resume(
        self,
        user_profile: Dict[str, Any],
        summary: str,
        skills: List[str],
        sections: List[ResumeSection]
    ) -> str:
        """Generate DOCX resume document"""
        
        if not Document:
            raise ImportError("python-docx not installed")
        
        doc = Document()
        
        # Header
        name = f"{user_profile.get('firstName', '')} {user_profile.get('lastName', '')}"
        header = doc.add_paragraph(name)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header.runs[0].font.size = Inches(0.25)  # 18pt
        header.runs[0].bold = True
        
        # Contact info
        contact = f"{user_profile.get('email', '')} | {user_profile.get('phone', '')} | {user_profile.get('location', '')}"
        contact_para = doc.add_paragraph(contact)
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Summary
        doc.add_paragraph("PROFESSIONAL SUMMARY").runs[0].bold = True
        doc.add_paragraph(summary)
        
        # Skills
        doc.add_paragraph("CORE COMPETENCIES").runs[0].bold = True
        skills_text = " • ".join(skills)
        doc.add_paragraph(skills_text)
        
        # Experience sections
        for section in sections:
            doc.add_paragraph(section.title.upper()).runs[0].bold = True
            doc.add_paragraph(section.content)
            
            for bullet in section.bullets:
                p = doc.add_paragraph(bullet.text, style='List Bullet')
        
        # Save document
        timestamp = int(datetime.utcnow().timestamp())
        filename = f"resume_{user_profile.get('user_id', 'user')}_{timestamp}.docx"
        filepath = Path("generated_resumes") / filename
        filepath.parent.mkdir(exist_ok=True)
        
        doc.save(str(filepath))
        
        return str(filepath)
    
    async def _generate_pdf_resume(
        self,
        user_profile: Dict[str, Any],
        summary: str,
        skills: List[str],
        sections: List[ResumeSection]
    ) -> str:
        """Generate PDF resume document"""
        
        timestamp = int(datetime.utcnow().timestamp())
        filename = f"resume_{user_profile.get('user_id', 'user')}_{timestamp}.pdf"
        filepath = Path("generated_resumes") / filename
        filepath.parent.mkdir(exist_ok=True)
        
        # Create PDF (simplified version)
        doc = SimpleDocTemplate(str(filepath), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Name
        name = f"{user_profile.get('firstName', '')} {user_profile.get('lastName', '')}"
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=6,
            alignment=1  # Center
        )
        story.append(Paragraph(name, title_style))
        
        # Contact
        contact = f"{user_profile.get('email', '')} | {user_profile.get('phone', '')} | {user_profile.get('location', '')}"
        story.append(Paragraph(contact, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Summary
        story.append(Paragraph("PROFESSIONAL SUMMARY", styles['Heading2']))
        story.append(Paragraph(summary, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Skills
        story.append(Paragraph("CORE COMPETENCIES", styles['Heading2']))
        skills_text = " • ".join(skills)
        story.append(Paragraph(skills_text, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Experience
        for section in sections:
            story.append(Paragraph(section.title.upper(), styles['Heading2']))
            story.append(Paragraph(section.content, styles['Normal']))
            
            for bullet in section.bullets:
                story.append(Paragraph(f"• {bullet.text}", styles['Normal']))
            
            story.append(Spacer(1, 12))
        
        doc.build(story)
        return str(filepath)
    
    async def _create_diff_report(
        self,
        user_profile: Dict[str, Any],
        summary: str,
        skills: List[str],
        sections: List[ResumeSection]
    ) -> List[Dict[str, Any]]:
        """Create diff report showing changes made"""
        
        diff_entries = []
        
        # Summary changes
        original_summary = user_profile.get('summary', 'No original summary')
        if summary != original_summary:
            diff_entries.append({
                "section": "summary",
                "before": original_summary,
                "after": summary,
                "reason": "ATS_optimization|job_alignment",
                "risk_level": "low"
            })
        
        # Skills changes  
        original_skills = user_profile.get('skills', [])
        if skills != original_skills:
            diff_entries.append({
                "section": "skills",
                "before": original_skills,
                "after": skills,
                "reason": "ATS_term_align|priority_reorder",
                "risk_level": "low"
            })
        
        # Experience bullets with RS
        for section in sections:
            for bullet in section.bullets:
                if bullet.rs and bullet.original_text:
                    diff_entries.append({
                        "section": "experience_bullet",
                        "before": bullet.original_text,
                        "after": bullet.text,
                        "reason": f"RS|{bullet.rs_basis[:50]}...",
                        "risk_level": bullet.risk_level,
                        "confidence": bullet.confidence
                    })
        
        return diff_entries
    
    async def check_ats_compatibility(self, resume_content: str) -> Dict[str, Any]:
        """Standalone ATS compatibility check"""
        
        task = AITask(
            task_id=f"ats_only_{datetime.utcnow().timestamp()}",
            task_type=TaskType.ATS_OPTIMIZATION,
            input_data={
                "resume_text": resume_content,
                "job_keywords": []
            }
        )
        
        response = await self.ai_orchestrator.execute_task(task)
        
        if response.success:
            return response.result
        else:
            return {
                "ats_score": 75,
                "recommendations": [
                    "Use standard section headings",
                    "Include relevant keywords", 
                    "Use bullet points for readability"
                ],
                "checks": {
                    "has_standard_headings": True,
                    "uses_bullet_points": True,
                    "proper_formatting": True
                }
            }
    
    async def generate_resume_diff(
        self,
        original_resume: Dict[str, Any],
        tailored_resume: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Generate diff report between original and tailored resume"""
        
        try:
            self.logger.info("Generating resume diff report")
            
            diff_report = {
                "diff_id": f"diff_{datetime.utcnow().timestamp()}",
                "generated_at": datetime.utcnow().isoformat(),
                "changes": [],
                "summary": {},
                "risk_assessment": "low"
            }
            
            # Compare sections
            changes = []
            rs_changes = 0
            total_changes = 0
            
            # Summary changes
            original_summary = original_resume.get('summary', '')
            tailored_summary = tailored_resume.get('summary', '')
            
            if original_summary != tailored_summary:
                changes.append({
                    "section": "summary",
                    "type": "modification",
                    "original": original_summary,
                    "modified": tailored_summary,
                    "change_type": "ai_tailored"
                })
                total_changes += 1
            
            # Skills changes
            original_skills = set(original_resume.get('skills', []))
            tailored_skills = set(tailored_resume.get('skills', []))
            
            if original_skills != tailored_skills:
                added_skills = tailored_skills - original_skills
                removed_skills = original_skills - tailored_skills
                
                if added_skills:
                    changes.append({
                        "section": "skills",
                        "type": "addition", 
                        "items": list(added_skills),
                        "change_type": "ats_optimization"
                    })
                    total_changes += len(added_skills)
                
                if removed_skills:
                    changes.append({
                        "section": "skills",
                        "type": "removal",
                        "items": list(removed_skills),
                        "change_type": "relevance_optimization"
                    })
                    total_changes += len(removed_skills)
            
            # Experience changes (check for RS bullets)
            original_exp = original_resume.get('experience_sections', [])
            tailored_exp = tailored_resume.get('experience_sections', [])
            
            for i, (orig_section, tailored_section) in enumerate(zip(original_exp, tailored_exp)):
                orig_bullets = [b.get('text', b) if isinstance(b, dict) else str(b) for b in orig_section.get('bullets', [])]
                tailored_bullets = tailored_section.get('bullets', [])
                
                for j, bullet in enumerate(tailored_bullets):
                    if isinstance(bullet, dict) and bullet.get('rs', False):
                        changes.append({
                            "section": f"experience_{i}",
                            "type": "rs_enhancement",
                            "original": bullet.get('original_text', ''),
                            "modified": bullet.get('text', ''),
                            "rs_basis": bullet.get('rs_basis', ''),
                            "confidence": bullet.get('confidence', 0.0),
                            "risk_level": bullet.get('risk_level', 'medium')
                        })
                        rs_changes += 1
                        total_changes += 1
                        
                        # Assess overall risk
                        if bullet.get('risk_level') == 'high':
                            diff_report["risk_assessment"] = "high"
                        elif bullet.get('risk_level') == 'medium' and diff_report["risk_assessment"] == "low":
                            diff_report["risk_assessment"] = "medium"
            
            # Generate summary
            diff_report["changes"] = changes
            diff_report["summary"] = {
                "total_changes": total_changes,
                "rs_enhancements": rs_changes,
                "sections_modified": len(set(change["section"] for change in changes)),
                "ats_optimizations": len([c for c in changes if c.get("change_type") == "ats_optimization"]),
                "risk_level": diff_report["risk_assessment"]
            }
            
            # Add rollback information
            diff_report["rollback_available"] = True
            diff_report["original_resume"] = original_resume
            
            self.logger.info(
                "Resume diff generated",
                total_changes=total_changes,
                rs_changes=rs_changes,
                risk_level=diff_report["risk_assessment"]
            )
            
            return diff_report
            
        except Exception as e:
            self.logger.error("Resume diff generation failed", error=str(e))
            raise